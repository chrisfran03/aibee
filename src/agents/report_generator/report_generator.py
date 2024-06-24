from jinja2 import Environment, BaseLoader
from docx import Document
from io import BytesIO
from src.llm import LLM

report_generator_prompt = open("src/agents/report_generator/prompt.jinja2").read().strip()

class ReportGenerator:
    def __init__(self, base_model, api_key):
        self.llm = LLM(base_model, api_key)

    def render(self, prompt):
        env = Environment(loader=BaseLoader())
        template = env.from_string(report_generator_prompt)

        return template.render(prompt=prompt)

    def validate_response(self, response: str) -> bool:
        return True
    
    # Parsing the response to provide suitable output 
    def parse_response(self, response: str):
        result = {"project": "", "focus": "", "plans": {}, "summary": ""}

        reply = ""
        current_section = None
        current_step = None

        for line in response.split("\n"):
            line = line.strip()

            if line.startswith("# Project Name:"):
                current_section = "project"
                result["project"] = line.split(":", 1)[1].strip()
            elif line.startswith("Your Reply to the Human Prompter:"):
                reply = line.split(":", 1)[1].strip()
            elif line.startswith("# Current Focus:"):
                current_section = "focus"
                result["focus"] = line.split(":", 1)[1].strip()
            elif line.startswith("# Plan:"):
                current_section = "plans"
            elif line.startswith("# Summary:"):
                current_section = "summary"
                result["summary"] = line.split(":", 1)[1].strip()
            elif current_section == "reply":
                result["reply"] += " " + line
            elif current_section == "focus":
                result["focus"] += " " + line
            elif current_section == "plans":
                if line.startswith("- [ ] Step"):
                    current_step = line.split(":")[0].strip().split(" ")[-1]
                    result["plans"][int(current_step)] = line.split(":", 1)[1].strip()
                elif current_step:
                    result["plans"][int(current_step)] += " " + line
            elif current_section == "summary":
                result["summary"] += " " + line.replace("```", "")

        result["project"] = result["project"].strip()
        result["focus"] = result["focus"].strip()
        result["summary"] = result["summary"].strip()

        return reply, result

    def execute(self, prompt: str) -> str:
        prompt = self.render(prompt)
        response = self.llm.inference(prompt)
        return response
    
    # def generate_report(self, response, filename):
    # # Check if filename ends with .docx, if not, append it
    #     if not filename.endswith('.docx'):
    #         filename += '.docx'
    
    #     doc = Document()
    #     lines = response.split('\n')

    #     for line in lines:
    #         if line.startswith('# '):  # Heading
    #             # Remove '# ' and add as a heading
    #             doc.add_heading(line[2:], level=1)
    #         elif line.startswith('## '):  # Subheading
    #             # Remove '## ' and add as a subheading
    #             doc.add_heading(line[3:], level=2)
    #         else:
    #             # Add regular paragraph
    #             doc.add_paragraph(line)

    #     doc.save(filename)
    #     print(f'DOCX file created: {filename}')
    def generate_report(self, response, filename):
        # Check if filename ends with .docx, if not, append it
        if not filename.endswith('.docx'):
            filename += '.docx'

        doc = Document()
        lines = response.split('\n')

        for line in lines:
            if line.startswith('# '):  # Heading
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):  # Subheading
                doc.add_heading(line[3:], level=2)
            else:
                doc.add_paragraph(line)

        # Save the document to an in-memory bytes buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)  # Rewind the buffer to start

        return buffer, filename

